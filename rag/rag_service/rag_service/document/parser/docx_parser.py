"""DOCX document parser using python-docx."""

from __future__ import annotations

import asyncio
import io
import logging
from functools import partial

from docx import Document

from rag_service.document.parser.base import DocumentParser, ParsedDocument
from rag_service.utils.errors import DocumentParseError

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Parser for DOCX files using python-docx."""

    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, partial(self._parse_sync, data, filename))

    def _parse_sync(self, data: bytes, filename: str) -> ParsedDocument:
        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise DocumentParseError(f"Failed to open DOCX '{filename}': {e}") from e

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        if not paragraphs:
            raise DocumentParseError(f"DOCX '{filename}' contains no extractable text")

        content = "\n\n".join(paragraphs)
        return ParsedDocument(
            content=content,
            metadata={"filename": filename, "parser": "docx"},
            title=filename,
        )
